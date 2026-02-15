"""
Document Processor Module for Knowledge Base System

This module handles ingestion of unstructured documents (PDF, DOCX, TXT) including:
- Text extraction from PDFs, DOCX, and TXT files
- Table extraction from PDFs using pdfplumber
- Text chunking using LangChain RecursiveCharacterTextSplitter
- Embedding generation using sentence-transformers
- Storage of chunks and embeddings in Supabase pgvector

Author: EDI.ai Team
Date: 2025-12-31
"""

import os
import logging
from typing import Dict, List, Optional, Tuple
import numpy as np
import pandas as pd
from sentence_transformers import SentenceTransformer
from sqlalchemy import create_engine

# Setup logging
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Process unstructured documents for knowledge bases.

    Features:
    - Extract text from PDFs, DOCX, TXT
    - Extract tables from PDFs using pdfplumber
    - Chunk text using LangChain RecursiveCharacterTextSplitter
    - Generate embeddings using sentence-transformers
    - Store chunks and embeddings in pgvector (Supabase)
    """

    def __init__(self, embedding_model: str = 'sentence-transformers/all-MiniLM-L6-v2'):
        """
        Initialize DocumentProcessor with embedding model.

        Args:
            embedding_model: Model name for sentence-transformers (default: all-MiniLM-L6-v2)
        """
        logger.info(f"Initializing DocumentProcessor with model: {embedding_model}")
        try:
            self.embedding_model = SentenceTransformer(embedding_model)
            self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
            logger.info(f"Embedding model loaded successfully. Dimension: {self.embedding_dim}")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
            raise

    def process_pdf(self, file_path: str, kb_id: str, doc_id: str) -> Dict:
        """
        Extract text and tables from PDF.

        Args:
            file_path: Path to PDF file
            kb_id: Knowledge base ID
            doc_id: Document ID from kb_documents table

        Returns:
            Dict containing:
                - text_chunks: List of text chunks
                - tables: List of extracted tables
                - page_count: Number of pages
                - has_tables: Boolean indicating if tables were found
        """
        import pdfplumber

        logger.info(f"Processing PDF: {file_path}")

        try:
            full_text = ""
            tables_data = []

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                logger.info(f"PDF has {page_count} pages")

                for page_num, page in enumerate(pdf.pages, start=1):
                    logger.debug(f"Processing page {page_num}/{page_count}")

                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        full_text += f"\n\n--- Page {page_num} ---\n\n{page_text}"

                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table and len(table) > 0:
                            try:
                                # Convert to pandas DataFrame
                                # First row is usually headers
                                headers = table[0]
                                data = table[1:]

                                # Clean headers (remove None, empty strings)
                                headers = [str(h) if h else f"Column_{i}" for i, h in enumerate(headers)]

                                df = pd.DataFrame(data, columns=headers)

                                tables_data.append({
                                    'page': page_num,
                                    'table_index': table_idx,
                                    'data': df.to_dict('records'),
                                    'columns': df.columns.tolist(),
                                    'row_count': len(df)
                                })

                                logger.info(f"Extracted table {table_idx} from page {page_num}: {len(df)} rows")
                            except Exception as e:
                                logger.warning(f"Failed to parse table {table_idx} on page {page_num}: {e}")

            # Chunk text
            chunks = self._chunk_text(full_text, chunk_size=1000, overlap=200)
            logger.info(f"Generated {len(chunks)} text chunks")

            return {
                'text_chunks': chunks,
                'tables': tables_data,
                'page_count': page_count,
                'has_tables': len(tables_data) > 0
            }

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

    def process_docx(self, file_path: str, kb_id: str, doc_id: str) -> Dict:
        """
        Extract text and tables from DOCX file.

        Args:
            file_path: Path to DOCX file
            kb_id: Knowledge base ID
            doc_id: Document ID

        Returns:
            Dict with text_chunks, tables, has_tables
        """
        from docx import Document

        logger.info(f"Processing DOCX: {file_path}")

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            logger.info(f"Extracted text from {len(doc.paragraphs)} paragraphs")

            # Extract tables
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                try:
                    # Get headers from first row
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    headers = [h if h else f"Column_{i}" for i, h in enumerate(headers)]

                    # Get data from remaining rows
                    data = []
                    for row in table.rows[1:]:
                        row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
                        data.append(row_data)

                    if data:  # Only add if table has data
                        df = pd.DataFrame(data)
                        tables_data.append({
                            'table_index': table_idx,
                            'data': df.to_dict('records'),
                            'columns': df.columns.tolist(),
                            'row_count': len(df)
                        })
                        logger.info(f"Extracted table {table_idx}: {len(df)} rows")

                except Exception as e:
                    logger.warning(f"Failed to parse table {table_idx}: {e}")

            # Chunk text
            chunks = self._chunk_text(full_text)
            logger.info(f"Generated {len(chunks)} text chunks")

            return {
                'text_chunks': chunks,
                'tables': tables_data,
                'has_tables': len(tables_data) > 0
            }

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    def process_txt(self, file_path: str, kb_id: str, doc_id: str) -> Dict:
        """
        Extract text from plain text file.

        Args:
            file_path: Path to TXT file
            kb_id: Knowledge base ID
            doc_id: Document ID

        Returns:
            Dict with text_chunks
        """
        logger.info(f"Processing TXT: {file_path}")

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                full_text = f.read()

            logger.info(f"Read {len(full_text)} characters from text file")

            # Chunk text
            chunks = self._chunk_text(full_text)
            logger.info(f"Generated {len(chunks)} text chunks")

            return {
                'text_chunks': chunks,
                'tables': [],
                'has_tables': False
            }

        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Chunk text using LangChain RecursiveCharacterTextSplitter.

        This uses the already-installed langchain-text-splitters package.

        Args:
            text: Full text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Overlap between chunks for context preservation

        Returns:
            List of text chunks
        """
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]
            )

            chunks = splitter.split_text(text)
            return chunks

        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            # Fallback to simple chunking
            logger.warning("Falling back to simple chunking")
            return self._simple_chunk(text, chunk_size, overlap)

    def _simple_chunk(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """Simple fallback chunking if LangChain fails."""
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap
        return chunks

    def generate_embeddings(self, texts: List[str]) -> np.ndarray:
        """
        Generate 384-dimensional embeddings using sentence-transformers.

        Args:
            texts: List of text chunks

        Returns:
            NumPy array of embeddings (n_texts, 384)
        """
        logger.info(f"Generating embeddings for {len(texts)} texts")

        try:
            embeddings = self.embedding_model.encode(
                texts,
                show_progress_bar=True,
                convert_to_numpy=True
            )
            logger.info(f"Generated embeddings shape: {embeddings.shape}")
            return embeddings

        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise

    def save_to_vector_db(self, kb_id: str, doc_id: str, chunks: List[str],
                          embeddings: np.ndarray, metadata: List[Dict],
                          supabase_client):
        """
        Batch insert chunks and embeddings to Supabase pgvector.

        Args:
            kb_id: Knowledge base ID
            doc_id: Document ID
            chunks: List of text chunks
            embeddings: NumPy array of embeddings
            metadata: List of metadata dicts (one per chunk)
            supabase_client: Supabase client instance
        """
        logger.info(f"Saving {len(chunks)} chunks to vector database")

        try:
            batch_data = []
            for idx, (chunk, embedding, meta) in enumerate(zip(chunks, embeddings, metadata)):
                batch_data.append({
                    'document_id': doc_id,
                    'kb_id': kb_id,
                    'chunk_index': idx,
                    'content': chunk,
                    'embedding': embedding.tolist(),  # Convert numpy to list for JSON
                    'chunk_metadata': meta
                })

            # Batch insert (Supabase handles batching efficiently)
            response = supabase_client.table('kb_document_chunks').insert(batch_data).execute()
            logger.info(f"Successfully saved {len(batch_data)} chunks to database")

            return response

        except Exception as e:
            logger.error(f"Error saving to vector database: {e}")
            raise


class TableExtractor:
    """
    Extract and process tables from documents for predictive analytics.

    This class handles:
    - Extracting specific tables from PDFs
    - Creating temporary SQLite databases for tables
    - Preparing tables for integration with PredictiveAnalyzer
    """

    def __init__(self):
        logger.info("Initializing TableExtractor")

    def extract_table_from_pdf(self, pdf_path: str, page_num: int, table_idx: int) -> Optional[pd.DataFrame]:
        """
        Extract specific table from PDF page.

        Args:
            pdf_path: Path to PDF file
            page_num: Page number (1-indexed)
            table_idx: Table index on that page (0-indexed)

        Returns:
            pandas DataFrame or None if table not found
        """
        import pdfplumber

        logger.info(f"Extracting table {table_idx} from page {page_num} of {pdf_path}")

        try:
            with pdfplumber.open(pdf_path) as pdf:
                if page_num < 1 or page_num > len(pdf.pages):
                    logger.error(f"Invalid page number: {page_num} (PDF has {len(pdf.pages)} pages)")
                    return None

                page = pdf.pages[page_num - 1]  # Convert to 0-indexed
                tables = page.extract_tables()

                if table_idx >= len(tables):
                    logger.error(f"Table index {table_idx} not found on page {page_num}")
                    return None

                table = tables[table_idx]

                # Convert to DataFrame
                headers = [str(h) if h else f"Column_{i}" for i, h in enumerate(table[0])]
                df = pd.DataFrame(table[1:], columns=headers)

                logger.info(f"Extracted table: {df.shape[0]} rows, {df.shape[1]} columns")
                return df

        except Exception as e:
            logger.error(f"Error extracting table: {e}")
            return None

    def create_temp_db_for_table(self, df: pd.DataFrame, kb_id: str, table_id: str) -> str:
        """
        Create temporary SQLite database for extracted table.

        This reuses the existing DataHandler pattern for compatibility with
        the PredictiveAnalyzer and SQL agents.

        Args:
            df: pandas DataFrame containing table data
            kb_id: Knowledge base ID
            table_id: Unique table identifier

        Returns:
            Path to created SQLite database
        """
        logger.info(f"Creating temp DB for table {table_id} in KB {kb_id}")

        try:
            # Create database path
            db_name = f"temp_db_kb_{kb_id}_table_{table_id}.db"
            temp_db_path = os.path.join("backend", db_name)

            # Create SQLite engine
            engine = create_engine(f'sqlite:///{temp_db_path}')

            # Write DataFrame to SQLite
            df.to_sql('data_table', engine, if_exists='replace', index=False)

            logger.info(f"Created temp database: {temp_db_path}")
            logger.info(f"Table 'data_table' has {len(df)} rows, {len(df.columns)} columns")

            return temp_db_path

        except Exception as e:
            logger.error(f"Error creating temp database: {e}")
            raise

    def validate_table_for_predictions(self, df: pd.DataFrame) -> Tuple[bool, str]:
        """
        Validate if table is suitable for predictive analytics.

        Args:
            df: pandas DataFrame

        Returns:
            Tuple of (is_valid, message)
        """
        # Check minimum rows
        if len(df) < 10:
            return False, f"Insufficient data: only {len(df)} rows (need at least 10)"

        # Check for numeric columns
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_cols) == 0:
            return False, "No numeric columns found for prediction"

        # Check for temporal columns (optional but helpful)
        temporal_patterns = ['date', 'time', 'year', 'month', 'day', 'period']
        temporal_cols = [col for col in df.columns
                        if any(pattern in col.lower() for pattern in temporal_patterns)]

        if len(temporal_cols) == 0:
            logger.warning("No temporal columns detected - predictions may be limited")

        return True, f"Table validated: {len(df)} rows, {len(numeric_cols)} numeric columns"


# Utility functions for integration with existing backend

def get_supabase_client(use_service_key=False):
    """
    Get Supabase client instance.

    This should import from your existing Supabase configuration.

    Args:
        use_service_key (bool): If True, forces use of service role key to bypass RLS
    """
    import os
    from supabase import create_client

    # Ensure environment variables are loaded (fallback)
    try:
        from dotenv import load_dotenv
        load_dotenv()
    except:
        pass

    url = os.getenv('NEXT_PUBLIC_SUPABASE_URL')

    if use_service_key:
        # Force use of service role key to bypass RLS for write operations
        key = os.getenv('SUPABASE_SERVICE_ROLE_KEY')
        if not key:
            raise ValueError("SUPABASE_SERVICE_ROLE_KEY not found in environment variables. This is required for knowledge base creation.")
    else:
        # Prefer service role key for backend operations to bypass RLS
        key = os.getenv('SUPABASE_SERVICE_ROLE_KEY') or os.getenv('NEXT_PUBLIC_SUPABASE_ANON_KEY')

    if not url or not key:
        raise ValueError("Supabase credentials not found in environment variables")

    return create_client(url, key)


# Example usage and testing
if __name__ == "__main__":
    # Setup logging for testing
    logging.basicConfig(level=logging.INFO)

    # Example: Process a PDF
    processor = DocumentProcessor()

    # Example PDF path (replace with actual file)
    test_pdf = "test_document.pdf"

    if os.path.exists(test_pdf):
        result = processor.process_pdf(test_pdf, "kb_123", "doc_456")
        print(f"Extracted {len(result['text_chunks'])} chunks")
        print(f"Found {len(result['tables'])} tables")

        # Generate embeddings
        if result['text_chunks']:
            embeddings = processor.generate_embeddings(result['text_chunks'][:5])
            print(f"Embeddings shape: {embeddings.shape}")
    else:
        print("Test PDF not found")
